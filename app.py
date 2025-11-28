import streamlit as st
import os
from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool
from docx import Document
import io

# =========================================================
# 1. KONFIGURASI HALAMAN
# =========================================================
st.set_page_config(page_title="Agen Material Pro", page_icon="🏗️")

# =========================================================
# 2. UI (TAMPILAN WEBSITE)
# =========================================================
st.title("🏗️ Agen Pengadaan Material Bangunan")
st.markdown("---")
st.write("Masukkan kebutuhan Anda, biarkan AI yang keliling toko online.")

col1, col2 = st.columns(2)
with col1:
    barang_input = st.text_input("Nama Barang", placeholder="Contoh: Semen Gresik 40kg")
with col2:
    lokasi_input = st.text_input("Lokasi", placeholder="Contoh: Jakarta Selatan")

tombol_cari = st.button("🚀 Mulai Pencarian")

# =========================================================
# 3. LOGIKA AGEN
# =========================================================
if tombol_cari:
    if not barang_input or not lokasi_input:
        st.warning("Mohon isi nama barang dan lokasi dulu ya, Bos!")
    else:
        with st.spinner('Sedang menghubungi Agen Hunter & Analyst... Mohon tunggu sebentar...'):
            try:
                # --- A. AMBIL KUNCI DARI BRANKAS CLOUD ---
                os.environ["SERPER_API_KEY"] = st.secrets["SERPER_API_KEY"]
                MY_GEMINI_KEY = st.secrets["GOOGLE_API_KEY"]
                
                # --- B. DEFINISI ALAT ---
                alat_search = SerperDevTool()

                # --- C. SETUP OTAK ---
                otak_gemini = LLM(
                    model="gemini/gemini-2.0-flash-001",
                    api_key=MY_GEMINI_KEY,
                    temperature=0.2 # Diturunkan biar lebih patuh/kaku
                )

                # --- D. DEFINISI AGEN ---
                hunter = Agent(
                    role='Material Hunter',
                    goal='Mencari link dan harga real-time',
                    backstory="Spesialis sourcing barang bangunan.",
                    tools=[alat_search],
                    llm=otak_gemini,
                    max_rpm=10
                )

                analyst = Agent(
                    role='Cost Analyst',
                    goal='Memilih penawaran terbaik',
                    backstory="Auditor harga yang teliti.",
                    llm=otak_gemini,
                    max_rpm=10
                )

                # --- E. DEFINISI TUGAS (UPDATED!) ---
                tugas_cari = Task(
                    description=f"Cari harga real-time '{barang_input}' di area '{lokasi_input}' via Tokopedia/Shopee. Kumpulkan URL LINK produknya secara lengkap.",
                    expected_output="List harga, nama toko, dan URL Link lengkap.",
                    agent=hunter
                )

                # DI SINI PERUBAHAN UTAMANYA:
                tugas_laporan = Task(
                    description="""
                    Buat laporan lengkap rekomendasi pembelian.
                    ATURAN PENTING:
                    1. Buat Tabel Perbandingan Harga.
                    2. Untuk rekomendasi terbaik, TULISKAN LINK (URL) SECARA LENGKAP DAN EKSPLISIT.
                    3. Jangan gunakan format [Link](url), tapi tulis urlnya langsung (contoh: https://tokopedia.com/...).
                    4. Pastikan Link tersebut ditaruh di baris baru agar mudah diklik/copy.
                    """,
                    expected_output="Laporan Markdown dengan URL Link yang terlihat jelas.",
                    agent=analyst
                )

                # --- F. EKSEKUSI ---
                tim = Crew(agents=[hunter, analyst], tasks=[tugas_cari, tugas_laporan])
                hasil = tim.kickoff()

                # --- G. TAMPILKAN HASIL ---
                st.success("Selesai! Berikut laporannya:")
                st.markdown("---")
                st.markdown(hasil)

                # =========================================================
                # FITUR DOWNLOAD WORD (.DOCX)
                # =========================================================
                doc = Document()
                doc.add_heading('Laporan Pengadaan Material', 0)
                
                # Kita tambahkan info barang yang dicari di judul
                doc.add_paragraph(f"Pencarian: {barang_input} - {lokasi_input}")
                doc.add_paragraph("-" * 50)
                
                doc.add_paragraph(str(hasil))
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.download_button(
                    label="📥 Download Laporan (.docx)",
                    data=bio.getvalue(),
                    file_name="laporan_belanja.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            except Exception as e:
                st.error(f"Terjadi kesalahan: {e}")

